# Copyright (c) 2024, AI Education Suite and contributors
# For license information, please see license.txt

"""
Backend for the custom 'AI Education Suite' dashboard page
(ai_core/page/ai_dashboard). Every method here is whitelisted and read-only
except download_question_paper_docx, which only reads + streams a file back
(no writes).
"""

import frappe

KPI_DEFS = [
	{
		"key": "high_risk_students",
		"label": "High-Risk Students",
		"doctype": "Student Risk Score",
		"filters": {"risk_level": "High", "status": "Open"},
		"icon": "\u26a0",
		"color": "#f97316",
	},
	{
		"key": "pending_grading",
		"label": "Pending Grading Suggestions",
		"doctype": "AI Grading Suggestion",
		"filters": {"status": "Pending Review"},
		"icon": "\u270d",
		"color": "#38bdf8",
	},
	{
		"key": "question_paper_drafts",
		"label": "Question Paper Drafts",
		"doctype": "AI Question Paper Draft",
		"filters": {"status": "Draft"},
		"icon": "\U0001F4C4",
		"color": "#a78bfa",
	},
	{
		"key": "house_allocations",
		"label": "Pending House Allocations",
		"doctype": "House Allocation Suggestion",
		"filters": {"status": "Pending"},
		"icon": "\u2302",
		"color": "#34d399",
	},
	{
		"key": "applicant_screenings",
		"label": "Applicant Screenings (Pending)",
		"doctype": "Applicant Screening Result",
		"filters": {"final_decision": "Pending"},
		"icon": "\U0001F393",
		"color": "#fbbf24",
	},
	{
		"key": "book_recommendations",
		"label": "Book Recommendations",
		"doctype": "Book Recommendation Log",
		"filters": {"status": "Suggested"},
		"icon": "\U0001F4DA",
		"color": "#f472b6",
	},
]

LIST_FIELDS = {
	"Student Risk Score": ["name", "student", "risk_score", "risk_level"],
	"AI Grading Suggestion": ["name", "student", "course", "suggested_score", "max_score"],
	"AI Question Paper Draft": ["name", "title", "course", "total_marks", "status"],
	"House Allocation Suggestion": ["name", "student", "suggested_house", "status"],
	"Applicant Screening Result": ["name", "applicant_name", "ai_score", "recommendation"],
	"Book Recommendation Log": ["name", "student", "status"],
}


def _get_kpi_def(key):
	kpi = next((k for k in KPI_DEFS if k["key"] == key), None)
	if not kpi:
		frappe.throw(f"Unknown KPI key: {key}")
	return kpi


@frappe.whitelist()
def get_kpis():
	result = []
	for kpi in KPI_DEFS:
		count = frappe.db.count(kpi["doctype"], kpi["filters"])
		result.append({**kpi, "value": count})
	return result


@frappe.whitelist()
def get_kpi_records(key):
	kpi = _get_kpi_def(key)
	fields = LIST_FIELDS.get(kpi["doctype"], ["name"])
	rows = frappe.get_list(
		kpi["doctype"],
		filters=kpi["filters"],
		fields=fields,
		order_by="modified desc",
		limit_page_length=15,
		ignore_permissions=True,
	)
	return {"doctype": kpi["doctype"], "label": kpi["label"], "rows": rows}


@frappe.whitelist()
def get_ai_insight():
	"""Fast, deterministic summary built from live counts. No API call."""
	kpis = get_kpis()
	lookup = {k["key"]: k["value"] for k in kpis}
	parts = []
	if lookup.get("high_risk_students"):
		parts.append(
			f"{lookup['high_risk_students']} student(s) are currently flagged high-risk and need a check-in."
		)
	if lookup.get("pending_grading"):
		parts.append(f"{lookup['pending_grading']} grading suggestion(s) are awaiting teacher review.")
	if lookup.get("question_paper_drafts"):
		parts.append(f"{lookup['question_paper_drafts']} question paper draft(s) are ready to finalize.")
	if lookup.get("applicant_screenings"):
		parts.append(f"{lookup['applicant_screenings']} applicant screening(s) are pending a decision.")
	if lookup.get("house_allocations"):
		parts.append(f"{lookup['house_allocations']} house allocation suggestion(s) haven't been actioned yet.")
	if not parts:
		return "Everything is caught up \u2014 no pending AI suggestions right now."
	return " ".join(parts)


@frappe.whitelist()
def get_ai_insight_llm():
	"""Regenerates the insight paragraph live via the configured LLM (Groq,
	through the shared claude_client wrapper) instead of the static template."""
	from ai_education_suite.ai_core.utils.claude_client import ClaudeClientError, call_claude

	kpis = get_kpis()
	summary_lines = "\n".join(f"- {k['label']}: {k['value']}" for k in kpis)
	prompt = (
		"You are summarizing a school admin dashboard for a busy principal. "
		"Given these current counts, write one short paragraph (2-3 sentences) "
		"highlighting what needs attention first and why. Be direct and "
		"specific, no preamble, no markdown formatting.\n\n" + summary_lines
	)
	try:
		return call_claude(prompt, max_tokens=200)
	except ClaudeClientError as e:
		frappe.throw(str(e))


@frappe.whitelist()
def get_question_paper_preview(name):
	draft = frappe.get_doc("AI Question Paper Draft", name)
	questions = [
		{
			"question_text": q.question_text,
			"topic": q.topic,
			"marks": q.marks,
			"difficulty": q.difficulty,
			"question_type": q.question_type,
		}
		for q in draft.questions
	]
	return {
		"title": draft.title,
		"course": draft.course,
		"student_group": draft.student_group,
		"total_marks": draft.total_marks,
		"questions": questions,
	}


@frappe.whitelist()
def download_question_paper_docx(name):
	try:
		from docx import Document
		from docx.enum.text import WD_ALIGN_PARAGRAPH
		from docx.shared import Inches, Pt
	except ImportError:
		frappe.throw(
			"python-docx is not installed on this bench. Run: "
			"./env/bin/pip install python-docx --break-system-packages "
			"then try again."
		)

	draft = frappe.get_doc("AI Question Paper Draft", name)

	doc = Document()
	section = doc.sections[0]
	section.left_margin = section.right_margin = Inches(1)

	title_p = doc.add_paragraph()
	title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	title_run = title_p.add_run(draft.title or "Question Paper")
	title_run.bold = True
	title_run.font.size = Pt(18)

	meta_bits = []
	if draft.course:
		meta_bits.append(f"Course: {draft.course}")
	if draft.student_group:
		meta_bits.append(f"Class: {draft.student_group}")
	if draft.total_marks:
		meta_bits.append(f"Total Marks: {draft.total_marks}")
	if meta_bits:
		meta_p = doc.add_paragraph()
		meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
		meta_run = meta_p.add_run("  |  ".join(meta_bits))
		meta_run.italic = True

	doc.add_paragraph()

	for i, q in enumerate(draft.questions, start=1):
		q_p = doc.add_paragraph()
		q_run = q_p.add_run(f"{i}. {q.question_text}")
		q_run.font.size = Pt(12)

		sub_p = doc.add_paragraph()
		sub_p.paragraph_format.left_indent = Inches(0.3)
		sub_run = sub_p.add_run(
			f"[{q.topic or ''} \u2014 {q.difficulty or ''} \u2014 {q.marks or 0} marks]"
		)
		sub_run.italic = True
		sub_run.font.size = Pt(9)

	from io import BytesIO

	buffer = BytesIO()
	doc.save(buffer)
	buffer.seek(0)

	filename = f"{(draft.title or 'question-paper').replace(' ', '_')}.docx"
	frappe.local.response.filename = filename
	frappe.local.response.filecontent = buffer.getvalue()
	frappe.local.response.type = "download"
